from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import _Cell
from docx.oxml.ns import qn
from PIL import ImageFont

def set_column_widths(table, col_widths):
    """
    Set fixed widths for table columns.
    
    Args:
        table: The table object from `python-docx`.
        col_widths: A list of widths in inches for each column.
    """
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(col_widths[idx])

def add_spacing(document, spacing=0):
    """
    Add vertical spacing using an empty paragraph with specified spacing size.

    Args:
        document: Document object.
        spacing: Size of spacing (in points)
    """
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = 0 
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = Pt(spacing)

def make_cell_cant_split(cell : _Cell) -> None:
    """
    Makes a table cell not break across pages. Prevents overflowing of table content.

    Args:
        cell: Table Cell object.
    """
    trPr = cell._tc.get_or_add_tcPr()
    trPr.append(OxmlElement('w:cantSplit'))
    

def add_content(document, image_path, text, image_width, text_width, font_size=12, font_family = 'Calibri', line_spacing = 1.5):
    """
    Add a centered table with an image on the left cell and text on the right cell.
    
    Args:
        document: Document object
        image_path: Path to the image file
        text: Text to display on the right
        font_size: Size of the text in points (12 by default for Intelife Easy Read Documents)
        font_family: Font family of the text (Calibri by default for Intelife Easy Read Documents)
    """
    # Add a table with one row and two cells
    table = document.add_table(rows=1, cols=2)
    # Adjust image cell and text cell width
    set_column_widths(table, [image_width, text_width])

    # Center the table
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.allow_autofit = False
    # Add image to left cell
    left_cell = table.cell(0, 0)
    left_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = left_cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(image_width-0.16))
    
    # Add text to right cell
    right_cell = table.cell(0, 1)
    right_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = right_cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_family
    
    make_cell_cant_split(left_cell)
    make_cell_cant_split(right_cell)

def measure_text_width(text, font_size, font_path = "test_fonts/calibri.ttf"):
    """
    Measures the width of specified text using Pillow.
    Currently uses normal Calibri as default.

    Args:
    text (str): Text content
    font_size (int): Font size in points.
    font_path (str): Path to the font file (optional).
    """
    font = ImageFont.truetype(font_path, font_size)
    # Measure the width of the entire text (in pixels)
    return font.getlength(text)

def estimate_row_height(text, image_size, font_size, column_width_in, line_spacing = 1.5, font_path = "test_fonts/calibri.ttf"):
    """
    Estimate the height of a row based on exact character widths using Pillow.
    
    Args:
        text (str): Text content in the row.
        image_size (float): Size of the image (in pixels).
        font_size (int): Font size in points.
        column_width_in (float): Width of the column in inches.
        line_spacing (float): Line spacing of document (1.5 for Easy Read Documents)
        font_path (str): Path to the font file (optional).
    
    Returns:
        The height of the box/table based on the image and text height.
    """
    # Convert column width to pixels
    column_width_pt = (column_width_in - 0.16) * 80 # 80 DPI (dots per inch), 0.08 inch is default side margin.

    words = text.split()
    lines = []
    current_line = ""
    for word in words:
        test_line = f"{current_line} {word}".strip()
        test_width = measure_text_width(test_line,font_size, font_path)
        if test_width <= column_width_pt:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
            current_line = word
    if current_line:
        lines.append(current_line)
    
    # Calculate height of text based on font size and line spacing.
    num_lines = len(lines)
    height_pt = (num_lines*(font_size * (line_spacing))) * 1.24 # 1.24 is factor for line spacing

    # Convert the calculated height to inch
    height_inch = height_pt/72
    return max(height_inch, image_size)

def create_page(doc, content_list, boxes_per_page, page_width, page_height, x_margin, y_margin):
    """
    Inserts up to `boxes_per_page` boxes into the word document based on margin calculations to fit within one page.

    Args:
        doc: Document object
        content_list: List of image+text tuples (default use case: LLM output)
        boxes_per_page: Maximum number of boxes to fit in the page
        page_width: Width of the page/document
        page_height: Height of the page/document
        x_margin: Total horizontal margins (space that cannot be occupied by content)
        y_margin: Total vertical margins (space that cannot be occupied by content)

    Returns:
        i: Number of groups successfully fitted into the page
    """
    # Calculate space available for content
    space_height_inch = page_height - y_margin
    space_width_inch = page_width - x_margin

    match boxes_per_page:
        case 3:
            image_width = 2 - 0.16 # 0.16 inch cell side margin
        case 4:
            image_width = 1.5 - 0.16 

    text_width = space_width_inch - image_width

    # Track the number of boxes and their combined height in the current page.
    i = 0
    total_height = 0

    for content in content_list:
        if i == boxes_per_page:
            break
        else:
            tmp = estimate_row_height(content['text'], image_width, 12, text_width)
            if (total_height + tmp > space_height_inch - 0.2 *(i+1)): # Check if the latest box can fit with a minimum margin of 0.2.
                break
            else:
                total_height += tmp
                i += 1

    spacing_inch = (space_height_inch-total_height) / (i+1)
    spacing_pt = spacing_inch * 72

    for content in content_list[0:i]:
        add_spacing(doc, spacing_pt)
        add_content(
            doc,
            content['image_path'],
            content['text'],
            image_width,
            text_width
        )
    return i

def create_document(output_path, content_list, boxes_per_page, header_text = "Header text", header_image = "test_images/logo.png", footer_text = "Footer Text", footer_image = "test_images/logo.png"):
    """
    Create a document with multiple bordered rectangles containing images and text, with specified header and footer.
    
    Args:
        output_path: Path where the document will be saved
        content_list: List of dictionaries containing:
            - image_path: Path to image file
            - text: Text to display
            - font_size: (optional) Font size in points
        boxes_per_page: Number of boxes per page desired
        header_text: Text of header
        header_image: Path of header image
        footer_text: Text of footer
        footer_image: Path of footer image
    """
    # Open an empty document to write to, top/bottom margin 1 inch, 
    doc = Document()
    
    # Add header (Use table for two cells, left for image, right for text.)
    doc.sections[0].header_distance = 0.2
    header = doc.sections[0].header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6))

    # Adjust header column widths and alignments
    set_column_widths(header_table, [0.5,5.5])
    header_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    header_table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Add header image
    header_logo_cell = header_table.cell(0, 0)
    header_logo_paragraph = header_logo_cell.paragraphs[0]
    header_logo_run = header_logo_paragraph.add_run()
    header_logo_run.add_picture(header_image, width=Inches(0.5-0.16))

    # Add header text
    header_text_cell = header_table.cell(0, 1)
    header_text_paragraph = header_text_cell.paragraphs[0]
    header_text_paragraph.text = header_text

    # Add footer
    doc.sections[0].footer_distance = 0.2
    footer = doc.sections[0].footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6))
    set_column_widths(footer_table, [5.5,0.5])
    footer_table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Adjust footer column widths and alignments
    footer_table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.TOP
    footer_table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add footer text 
    footer_text_cell = footer_table.cell(0,0)
    footer_text_paragraph = footer_text_cell.paragraphs[0]
    footer_text_paragraph.text = footer_text
    footer_text_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Add footer image
    footer_logo_cell = footer_table.cell(0,1)
    footer_logo_paragraph = footer_logo_cell.paragraphs[0]
    footer_logo_run = footer_logo_paragraph.add_run()
    footer_logo_run.add_picture(footer_image, width = Inches(0.5-0.16))
    
    # Obtain page dimensions and margins for calculation.
    section = doc.sections[0]
    page_height, page_width = section.page_height.inches, section.page_width.inches
    x_margin = section.left_margin.inches + section.right_margin.inches
    y_margin = section.top_margin.inches + section.bottom_margin.inches

    i =  0
    while i < len(content_list):
        n = create_page(doc, content_list[i:i+boxes_per_page], boxes_per_page, page_width, page_height, x_margin, y_margin)
        # Advance the group index by the number of groups placed on this page
        i += n
        if (i < len(content_list)):
            doc.add_page_break()

    # Save the document
    doc.save(output_path)

def create_document_template(output_path, content_list, boxes_per_page, template_path = 'intelife-template.docx'):
    """
    Create a document with multiple bordered rectangles containing images and text, with an existing .docx document for the formatting.
    
    Args:
        output_path: Path where the document will be saved
        content_list: List of dictionaries containing:
            - image_path: Path to image file
            - text: Text to display
            - font_size: (optional) Font size in points
        boxes_per_page: Number of boxes per page desired
        header_text: Text of header
        header_image: Path of header image
        footer_text: Text of footer
        footer_image: Path of footer image
    """
    doc = Document(template_path)

    if doc.paragraphs and not doc.paragraphs[0].text.strip():
        # Remove the first paragraph if it's blank
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
        p._element = None

    section = doc.sections[0]
    page_height, page_width = section.page_height.inches, section.page_width.inches
    x_margin = section.left_margin.inches + section.right_margin.inches
    y_margin = section.top_margin.inches + section.bottom_margin.inches + doc.sections[0].header_distance.inches + doc.sections[0].footer_distance.inches

    i =  0
    while i < len(content_list):
        n = create_page(doc, content_list[i:i+boxes_per_page], boxes_per_page, page_width, page_height, x_margin, y_margin)
        # Advance the group index by the number of groups placed on this page
        i += n
        if (i < len(content_list)):
            doc.add_page_break()

    # Save the document
    doc.save(output_path)


# TEST CASE
''' 
content_list = [
    {
        'image_path': 'test_images/Standard1.png',
        'text': 'First item description'
    },
    {
        'image_path': 'test_images/Standard2.png',
        'text': 'Second item description. Second item description. Second item description. Second item description. Second item description.'
    },
    {
        'image_path': 'test_images/Standard3.png',
        'text': 'Third item descriptionSecond item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item descript. item descSecond item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item description. Second item descript. item descripripripriprppppppppppppppppppppp ddddddddd'
    },
    {
        'image_path': 'test_images/Standard1.png',
        'text': 'Fourth item description'
    },
    {
        'image_path': 'test_images/Standard1.png',
        'text': 'Fifth item description'
    },
        {
        'image_path': 'test_images/Standard1.png',
        'text': 'Sixth item description'
    },
        {
        'image_path': 'test_images/Standard1.png',
        'text': 'Seventh item description'
    }
]


create_document('output.docx', content_list, 4)
create_document_template('output_temp.docx', content_list, 4)
'''